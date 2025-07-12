import PyPDF2
import docx
import logging
from typing import List, Dict, Any
from datetime import datetime
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# Initialize logger
logger = logging.getLogger("aviator_chatbot")


class DocumentProcessor:
    """Handle document loading and processing."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        # Updated separators for RecursiveCharacterTextSplitter for better chunking
        # It prioritizes larger, more meaningful breaks like multiple newlines
        # and then progressively smaller ones. Markdown-like headings are included
        # in case your extracted text contains them.
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=[
                "\n\n\n",  # Triple newline for major section breaks
                "\n\n",  # Double newline for paragraph breaks
                "\n",  # Single newline for line breaks
                "---",  # Common horizontal rule
                "***",  # Common horizontal rule
                "## ",  # Markdown H2 (if present in plain text after extraction)
                "### ",  # Markdown H3
                "#### ",  # Markdown H4
                "##### ",  # Markdown H5
                "###### ",  # Markdown H6
                ". ",  # Sentence end (with space)
                "! ",  # Sentence end (with space)
                "? ",  # Sentence end (with space)
                ", ",  # Comma (less ideal, but fallback)
                " ",  # Space (word split)
                "",  # Character by character (last resort)
            ],
            is_separator_regex=False
        )
        logger.info(f"DocumentProcessor initialized with chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")

    def load_document(self, file_path: str, file_type: str) -> List[str]:
        """Load document based on file type."""
        logger.info(f"Loading document: {file_path} (Type: {file_type})")

        # --- FIX: Ensure extracted content is always valid before returning ---
        raw_text_content = []
        if file_type == "pdf":
            raw_text_content = self._load_pdf(file_path)
        elif file_type == "docx":
            raw_text_content = self._load_docx(file_path)
        elif file_type == "txt":
            raw_text_content = self._load_txt(file_path)
        else:
            logger.error(f"Unsupported file type: {file_type}")
            raise ValueError(f"Unsupported file type: {file_type}")

        # Filter out any empty or whitespace-only strings from the loaded content
        valid_text_content = [t for t in raw_text_content if t and t.strip()]
        if not valid_text_content:
            logger.warning(f"No valid text content extracted from {file_path}. Returning empty list.")
        return valid_text_content

    def _load_pdf(self, file_path: str) -> List[str]:
        """Load PDF document."""
        try:
            full_text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    # --- FIX: Only add text if it's not empty or just whitespace ---
                    if page_text and page_text.strip():
                        # Optionally add page numbers for context, even in this simple setup
                        # text += f"--- Page {page_num + 1} ---\n{page_text.strip()}\n\n"
                        full_text += page_text.strip() + "\n"  # Continue with original style, just stripped
                logger.info(f"Successfully loaded PDF: {file_path}")
            # --- FIX: Return empty list if no meaningful text was extracted ---
            return [full_text] if full_text.strip() else []
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise

    def _load_docx(self, file_path: str) -> List[str]:
        """Load DOCX document."""
        try:
            doc = docx.Document(file_path)
            full_text = ""
            for paragraph in doc.paragraphs:
                # --- FIX: Only add text if it's not empty or just whitespace ---
                if paragraph.text and paragraph.text.strip():
                    full_text += paragraph.text.strip() + "\n"
            logger.info(f"Successfully loaded DOCX: {file_path}")
            # --- FIX: Return empty list if no meaningful text was extracted ---
            return [full_text] if full_text.strip() else []
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise

    def _load_txt(self, file_path: str) -> List[str]:
        """Load TXT document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            logger.info(f"Successfully loaded TXT: {file_path}")
            # --- FIX: Return empty list if no meaningful text was extracted ---
            return [text] if text.strip() else []
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            raise

    def chunk_documents(self, texts: List[str], metadata: Dict[str, Any]) -> List[Document]:
        """Chunk documents with metadata."""
        documents = []
        for text_idx, text in enumerate(texts):
            if not text or not text.strip():
                logger.warning(f"Skipping empty or whitespace-only text input for chunking (index: {text_idx}).")
                continue
            chunks = self.text_splitter.split_text(text)
            logger.info(f"Chunked text {text_idx + 1} into {len(chunks)} pieces.")
            for i, chunk in enumerate(chunks):
                if not chunk or not chunk.strip():
                    logger.debug(f"Skipping empty or whitespace-only chunk (index: {i}) from text {text_idx + 1}.")
                    continue
                doc_metadata = metadata.copy()
                doc_metadata['chunk_id'] = i
                doc_metadata['timestamp'] = datetime.now().isoformat()
                documents.append(Document(page_content=chunk.strip(), metadata=doc_metadata))
        logger.info(f"Total {len(documents)} document chunks created with metadata.")
        return documents