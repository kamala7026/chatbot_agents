import streamlit as st
from streamlit_modal import Modal
import os
import uuid
from typing import List, Dict, Any
import tempfile
from datetime import datetime

from langchain_core.language_models import BaseChatModel
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import PyPDF2
import docx
from tavily import TavilyClient
from langchain.vectorstores import Chroma
from langchain.embeddings import HuggingFaceEmbeddings

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.agents import Tool, AgentExecutor, create_react_agent
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferWindowMemory
from langchain.schema import Document as LangchainDocument

from core.document_management import DocumentManager
from ticket_create_agent import TicketCreationTool

# --- UI Customization - NEW SECTION ---
# Inject custom CSS for basic styling and the loading dots animation
st.markdown("""
<style>
/* General body and main content adjustments */
.main {
    padding-top: 20px; /* Adjust top padding */
    padding-left: 20px;
    padding-right: 20px;
    padding-bottom: 20px;
}

/* Specific styling for chat messages for a more compact look */
.stChatMessage {
    margin-bottom: 5px; /* Reduce space between messages */
    padding: 5px 10px; /* Reduce internal padding of message bubbles */
    border-radius: 10px; /* Slightly more rounded corners */
}

.stChatMessage [data-testid="chatMessageUser"] {
    background-color: #e0f2f7; /* Light blue for user messages */
    text-align: right; /* Align user message content to right */
    border-bottom-right-radius: 0; /* Make bottom-right corner sharp */
}

.stChatMessage [data-testid="chatMessageAssistant"] {
    background-color: #f0f2f6; /* Light grey for assistant messages */
    text-align: left; /* Align assistant message content to left */
    border-bottom-left-radius: 0; /* Make bottom-left corner sharp */
}

/* Avatar styling */
.stChatMessage [data-testid="stChatMessageAvatar"] {
    width: 30px; /* Smaller avatar size */
    height: 30px;
    font-size: 18px; /* Adjust font size for emojis */
    line-height: 30px; /* Center emoji vertically */
    border-radius: 50%; /* Make avatar perfectly round */
    background-color: #cce0ff; /* Light blue background for avatars */
    display: flex;
    justify-content: center;
    align-items: center;
    color: #333; /* Darker text for emojis */
}

/* Reduce padding around the chat input area */
.stChatInputContainer {
    padding: 10px 0;
}

/* Hide Streamlit header, footer, and default sidebar toggle for a cleaner look */
header.st-emotion-cache-zt5ig8 { /* Specific class for the top header */
    visibility: hidden;
    height: 0%;
}
.st-emotion-cache-cio0dv { /* Specific class for the main menu button */
    visibility: hidden;
}
.st-emotion-cache-ysnmbv { /* Specific class for footer */
    visibility: hidden;
}
/* Adjust main content padding when sidebar is hidden */
.block-container {
    padding-top: 0rem;
    padding-bottom: 0rem;
    padding-left: 0rem;
    padding-right: 0rem;
}

/* Target the main app container to constrain its width */
.css-1r6dm1u { /* You might need to inspect this class name in your browser if it changes */
    max-width: 800px; /* Set a max width for the chat window, adjust as needed */
    margin-left: auto;
    margin-right: auto;
    border: 1px solid #ddd; /* Add a subtle border */
    box-shadow: 0 4px 8px rgba(0,0,0,0.1); /* Add a subtle shadow */
    border-radius: 10px; /* Rounded corners for the entire chat window */
    overflow: hidden; /* Hide overflow content from border-radius */
}

/* Sidebar styling for custom header */
[data-testid="stSidebarContent"] {
    padding-top: 20px; /* Add some padding at the top of the sidebar */
}

/* Adjust the chat input background */
div.st-emotion-cache-1c7y2kl:has(div.st-emotion-cache-fybftg) {
    background-color: #f0f2f6; /* Lighter background for the input area */
    border-radius: 15px; /* Rounded input field */
    border: 1px solid #ddd;
    padding: 8px 15px; /* Padding inside the input container */
}

/* Specific targeting for the text input within the chat input */
div.st-emotion-cache-1c7y2kl input[data-testid="stTextInput"] {
    background-color: transparent !important; /* Transparent background for the actual text input */
    border: none !important; /* Remove border */
    box-shadow: none !important; /* Remove shadow */
    padding: 0; /* Remove internal padding if any */
}

/* --- Loading Dots Animation --- */
.loading-dots-container {
    display: flex;
    align-items: center;
    justify-content: flex-start; /* Align dots to the left */
    padding: 5px 0; /* Adjust padding as needed */
    min-height: 25px; /* Ensure space for dots */
}

.loading-dot {
    width: 8px;
    height: 8px;
    background-color: #6495ED; /* Color of the dots */
    border-radius: 50%;
    margin: 0 3px; /* Space between dots */
    animation: bounce 1.4s infinite ease-in-out both;
}

.loading-dot:nth-child(1) { animation-delay: -0.32s; }
.loading-dot:nth-child(2) { animation-delay: -0.16s; }
.loading-dot:nth-child(3) { animation-delay: 0s; }

@keyframes bounce {
    0%, 80%, 100% { transform: scale(0); }
    40% { transform: scale(1); }
}

</style>
""", unsafe_allow_html=True)


# Configure Streamlit
st.set_page_config(
    page_title="Aviator Chatbot", # Changed page title
    page_icon="✈️", # Changed page icon to match "Aviator" theme
    layout="wide",
    initial_sidebar_state="expanded" # Keep sidebar expanded initially for settings
)

# Constants
CATEGORIES = ["TGO", "LENS", "AO", "AIC"]
STATUS_OPTIONS = ["Active", "Inactive"]
ACCESS_OPTIONS = ["Internal", "External"]
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
CHROMA_PERSIST_DIR = "./chroma_db_docs"

# --- Global Configuration / Caching ---
# Disable ChromaDB telemetry globally (recommended placement)
os.environ['CHROMA_ANALYTICS'] = 'False'
os.environ['TAVILY_API_KEY'] = "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"
# Use st.cache_resource for the embeddings model to load it only once
@st.cache_resource(show_spinner=False)
def load_embeddings_model_cached():
    """Loads and caches the HuggingFace embeddings model."""
    import torch
    device = "cuda" if torch.cuda.is_available() else "cpu"
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': device}
    )
    return embeddings

# --- DocumentProcessor, RAGRetriever, InternetSearchTool, RAGChatbot Classes remain UNCHANGED ---
# You would keep these classes as they are from your previous code.
# I'm omitting them here for brevity but assume they are present and correct.
# ... (your existing DocumentProcessor, RAGRetriever, InternetSearchTool, RAGChatbot classes) ...

class DocumentProcessor:
    """Handle document loading and processing"""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )

    def load_document(self, file_path: str, file_type: str) -> List[str]:
        """Load document based on file type"""
        if file_type == "pdf":
            return self._load_pdf(file_path)
        elif file_type == "docx":
            return self._load_docx(file_path)
        elif file_type == "txt":
            return self._load_txt(file_path)
        else:
            # Re-raise ValueError to be caught by st.status in add_document
            raise ValueError(f"Unsupported file type: {file_type}")

    def _load_pdf(self, file_path: str) -> List[str]:
        """Load PDF document"""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return [text]

    def _load_docx(self, file_path: str) -> List[str]:
        """Load DOCX document"""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return [text]

    def _load_txt(self, file_path: str) -> List[str]:
        """Load TXT document"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return [file.read()]

    def chunk_documents(self, texts: List[str], metadata: Dict[str, Any]) -> List[Document]:
        """Chunk documents with metadata"""
        documents = []
        for text in texts:
            chunks = self.text_splitter.split_text(text)
            for i, chunk in enumerate(chunks):
                doc_metadata = metadata.copy()
                doc_metadata['chunk_id'] = i
                doc_metadata['timestamp'] = datetime.now().isoformat()
                documents.append(Document(page_content=chunk, metadata=doc_metadata))
        return documents



class RAGRetriever:
    """Simple RAG retriever with metadata filtering and multi-query."""

    def __init__(self, vectorstore, llm: BaseChatModel, user_type: str = "non-support"):
        self.vectorstore = vectorstore
        self.user_type = user_type
        self.llm = llm # Store the LLM

        # Define the query generation prompt - moved to __init__ for single definition
        self.query_gen_prompt = PromptTemplate.from_template("""
        You are a helpful AI assistant. Your task is to generate 3 different versions of the given user question that are semantically similar but use different phrasing.
        These variations will be used to retrieve more comprehensive results from a document database.

        Original question: {question}

        Generated questions (one per line, do not number them):
        """)

    def get_relevant_documents(self, query: str) -> List[LangchainDocument]:
        """Retrieve relevant documents with metadata filtering and multi-query."""
        query_variations = self._generate_query_variations(query)
        st.sidebar.info(f"Generated query variations: {query_variations[:2]}...") # Display first 2 for brevity in sidebar

        all_results = []
        seen_content = set() # Use a set to keep track of unique document content

        for q in query_variations:
            try:
                results = self.vectorstore.similarity_search(
                    q,
                    k=5 # Retrieve more potential candidates per query
                )
                filtered_results = self._filter_by_metadata(results)

                for doc in filtered_results:
                    # Add only unique documents based on content
                    if doc.page_content not in seen_content:
                        all_results.append(doc)
                        seen_content.add(doc.page_content)
            except Exception as e:
                print(f"Error in similarity search for '{q}': {e}") # Use print for console visibility
                continue

        # Return top 5 unique results from the combined set
        return all_results[:5]

    def _generate_query_variations(self, original_query: str) -> List[str]:
        """Generate query variations for multi-query retrieval using the LLM."""
        try:
            query_generator_chain = self.query_gen_prompt | self.llm.bind(stop=["\n\n"]) | (lambda x: x.content.split('\n'))
            variations = query_generator_chain.invoke({"question": original_query})
            variations = [v.strip() for v in variations if v.strip()]
            if original_query not in variations:
                variations.insert(0, original_query)
            return variations
        except Exception as e:
            st.warning(f"Error generating query variations: {e}. Falling back to original query only.")
            return [original_query]

    def _filter_by_metadata(self, documents: List[LangchainDocument]) -> List[LangchainDocument]:
        """Filter documents based on metadata"""
        filtered = []
        for doc in documents:
            metadata = doc.metadata
            if metadata.get("status", "").lower() != "active":
                continue
            if self.user_type == "non-support":
                if metadata.get("access", "").lower() != "external":
                    continue
            filtered.append(doc)
        return filtered

    # Removed _remove_duplicates as it's integrated into get_relevant_documents



class InternetSearchTool:
    """Tool for internet search using Tavily API."""

    def __init__(self):
        # Retrieve Tavily API key from environment variables
        self.TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")

        if not self.TAVILY_API_KEY:
            raise ValueError("Tavily API Key (TAVILY_API_KEY) not found in environment variables. Get it from tavily.com")

        self.tavily_client = TavilyClient(api_key=self.TAVILY_API_KEY)

    def search(self, query: str, num_results: int = 3) -> str:
        """
        Search the internet using Tavily API.
        Returns a concise answer and top N relevant source links.
        """
        try:
            # Use Tavily's search method
            # You can set include_answer=True to get a summarized answer directly from Tavily
            # and include_raw_content=False to save on token usage if you only need the answer/sources.
            response = self.tavily_client.search(
                query=query,
                search_depth="advanced", # "basic" or "advanced" (more comprehensive, but slower/costlier)
                max_results=num_results,
                include_answer=True,    # Get a summarized answer
                include_raw_content=False, # Don't include raw content of pages unless specifically needed by LLM
            )

            # Extract summarized answer if available
            answer = response.get('answer', 'No concise answer found.')

            # Extract source documents
            sources = response.get('results', [])
            formatted_sources = []
            if sources:
                for i, source in enumerate(sources):
                    title = source.get('title', 'N/A')
                    url = source.get('url', 'N/A')
                    # snippet = source.get('content', 'No snippet available.') # Content is usually raw, maybe too long.
                    formatted_sources.append(f"{i + 1}. {title}\n   URL: {url}\n")
            else:
                formatted_sources.append("No relevant sources found.")

            # Combine answer and sources for the LLM
            output = f"Tavily Answer: {answer}\n\nSources:\n" + "\n".join(formatted_sources)

            return output

        except Exception as e:
            # Handle API-specific errors or general network issues
            return f"Error performing internet search with Tavily: {str(e)}"


class RAGChatbot:
    """Main RAG Chatbot class"""

    def __init__(self):
        self.embeddings = None
        self.vectorstore = None
        self.llm = None
        self.document_processor = DocumentProcessor()
        self.internet_search = InternetSearchTool()
        self.memory = ConversationBufferWindowMemory(
            memory_key="chat_history",
            return_messages=True,
            k=5
        )
        self.agent_executor = None

    def initialize_components(self, google_api_key: str):
        """Initialize LLM, embeddings, and vector store"""
        try:
            # Use st.status for a consolidated initialization progress display
            with st.status("Initializing Chatbot Components...", expanded=True, state="running") as status:
                self.embeddings = load_embeddings_model_cached() # Use the cached function
                if self.embeddings:
                    status.write("✅ Embeddings model loaded.")


                #status.update(label="Initializing LLM...", state="running",expanded=True)
                self.llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    google_api_key=google_api_key,
                    temperature=0.7
                )
                status.write("✅ LLM initialized.")

                #status.update("Setting up vector database (ChromaDB)...", state="running")
                if not os.path.exists(CHROMA_PERSIST_DIR):
                    os.makedirs(CHROMA_PERSIST_DIR)
                    status.write(f"Created directory: {CHROMA_PERSIST_DIR}")

                try:
                    self.vectorstore = Chroma(
                        persist_directory=CHROMA_PERSIST_DIR,
                        embedding_function=self.embeddings,
                        collection_name="documents"
                    )
                    count = self.vectorstore._collection.count()
                    status.write(f"Vector store initialized. Current document count: {count}")

                except Exception as chroma_error:
                    status.warning(f"ChromaDB direct initialization failed: {chroma_error}. Trying persistent client...")
                    # Fallback to direct PersistentClient instantiation
                    import chromadb
                    client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)
                    # get_or_create_collection is safer
                    collection = client.get_or_create_collection("documents", embedding_function=self.embeddings) # ensure embedding_function is passed here
                    self.vectorstore = Chroma(
                        client=client,
                        collection_name="documents",
                        embedding_function=self.embeddings # Still pass here for LangChain wrapper
                    )
                    count = self.vectorstore._collection.count()
                    status.write(f"Alternative ChromaDB init successful. Current document count: {count}")

                if self.vectorstore is None:
                    raise Exception("Vector store failed to initialize after all attempts.")

                status.update(label=f"Chatbot Ready with {count} document chunks!", state="complete", expanded=False)
                return True

        except Exception as e:
            status.error(f"❌ Error during chatbot initialization: {str(e)}")
            status.update(label="Chatbot Initialization Failed!", state="error", expanded=True)
            self.embeddings = None
            self.llm = None
            self.vectorstore = None
            return False

    def setup_agent(self, user_type: str = "non-support"):
        """Setup LangChain agent with tools"""
        retriever = RAGRetriever(self.vectorstore, self.llm, user_type)

        def get_current_chat_history() -> List[Any]:
            return self.memory.load_memory_variables({})['chat_history']
        ticket_tool_instance = TicketCreationTool(self.llm, chat_history_provider=get_current_chat_history)

        def rag_search(query: str) -> str:
            """Search for information in uploaded documents."""
            try:
                docs = retriever.get_relevant_documents(query)
                if not docs:
                    return "No relevant documents found in the database."
                context = "\n\n".join([doc.page_content for doc in docs])
                source_info = "\n".join([
                    f"Source: {doc.metadata.get('filename', 'Unknown')} (Category: {doc.metadata.get('category', 'Unknown')})"
                    for doc in docs
                ])
                return f"Context from documents:\n{context}\n\nSources:\n{source_info}"
            except Exception as e:
                return f"Error searching documents: {str(e)}"

        def perform_internet_search(query: str) -> str:
            """Search the internet for general knowledge."""
            return self.internet_search.search(query)

        def default_response(query: str) -> str:
            """Handles general greetings or casual conversation."""
            return "Hi there! How can I help you today?"

        tools = [
            Tool(
                name="rag_retriever",
                func=rag_search,
                description="Search for information in uploaded documents. Use this first for any query that is likely to be answered by documents."
            ),
            Tool(
                name="internet_search",
                func=perform_internet_search,
                description="Search the internet when information is not found in documents or for general knowledge questions."
            ),
            Tool(
                name="default_response",
                func=default_response,
                description="Use this when the user says hello, hi, thank you, or any casual greeting or closing. It provides a direct friendly response."
            ),
            ticket_tool_instance.get_tool()
        ]

        prompt = PromptTemplate.from_template("""
                You are an AI assistant that helps users by answering questions based on documents or by searching the internet.
                You use a ReAct agent with the following tools:

                - rag_retriever: Use this tool to search uploaded documents for specific information.
                - internet_search: Use this tool to search the web for information not found in documents or for general knowledge.
                - default_response: Use this tool for greetings or casual conversation (e.g., "hi", "hello", "thank you"). When this tool is used, its output is often the final answer.
                - create_support_ticket: Use this tool when the user explicitly asks to 'create a ticket', 'raise an issue', 'open a support request', or similar. This tool will manage the collection of required details.

                IMPORTANT RULES:
                - Always respond using the ReAct format:
                  Thought: Your reasoning here.
                  Action: The name of the tool to use.
                  Action Input: The input to the tool (a string).
                  OR
                  Thought: Your reasoning here, leading to a final answer.
                  Final Answer: Your final answer to the user.

                - If the user's input is a simple greeting (like "hi", "hello"), a thank you, or other casual conversational remark, **you must use the `default_response` tool**. The output of this tool will typically be your `Final Answer`.
                - For any factual or information-seeking question, prioritize `rag_retriever` first, then `internet_search`. Don't sumerize the content if the content is structured data like list of points or tabular data
                - If the user expresses an intent to create a ticket, regardless of whether they provide details upfront, you must use the `create_support_ticket` tool. This tool will then handle extracting information and asking for missing details.
                - **When a tool (especially `create_support_ticket` or `default_response`) provides a complete, user-facing answer, your next step should be `Final Answer: <Tool Output>`. Do NOT attempt to take another Action if the tool's output is meant to be the final response to the user's current request.**
                - Do not generate `Final Answer:` simultaneously with `Action:` in the same turn. If you decide to take an `Action:`, the `Final Answer:` should come in a subsequent turn after the tool's output is received.
                - If you find the answer from Documents then try to provide same format as that mentioned in document always otherwise sumarise it for final answer.
                EXAMPLES:

                User: Hello!
                Thought: The user greeted me. I should use the default_response tool for a friendly greeting.
                Action: default_response
                Action Input: hello
                Observation: Hi there! How can I help you today?
                Final Answer: Hi there! How can I help you today?

                User: What is the AIC policy?
                Thought: The user is asking a question about a specific policy, which is likely in the documents. I should use the rag_retriever tool.
                Action: rag_retriever
                Action Input: What is the AIC policy?
                Observation: Context from documents: ... (document content) ...
                Final Answer: Based on the AIC policy document, ... (summarized answer) ...

                User: Search for the latest TGO guidelines
                Thought: The user is asking for current guidelines, which might require external search. I will use the internet_search tool.
                Action: internet_search
                Action Input: latest TGO guidelines
                Observation: Internet search results for 'latest TGO guidelines': ... (search results) ...
                Final Answer: I found the following information about the latest TGO guidelines: ... (summarize search results) ...

                User: I want to create a ticket.
                Thought: The user explicitly stated intent to create a ticket. I should use the create_support_ticket tool to begin the process of gathering details.
                Action: create_support_ticket
                Action Input: I want to create a ticket.
                Observation: To create the ticket, I need a few more details: department (IT, HR, or PS), issue severity (High, Medium, or Low), impacted client name, impacted time (e.g., 'yesterday', 'today 2 PM'). Please provide them.
                Final Answer: To create the ticket, I need a few more details: department (IT, HR, or PS), issue severity (High, Medium, or Low), impacted client name, impacted time (e.g., 'yesterday', 'today 2 PM'). Please provide them.

                User: Can you raise an issue? It's for HR, high severity, for John Doe, and it happened yesterday.
                Thought: The user wants to raise an issue and provided all necessary details. I should use the create_support_ticket tool to create the ticket.
                Action: create_support_ticket
                Action Input: Can you raise an issue? It's for HR, high severity, for John Doe, and it happened yesterday.
                Observation: Ticket **TICKET-XXXX** created successfully with the following details: ... Is there anything else I can help you with?
                Final Answer: Ticket **TICKET-XXXX** created successfully with the following details: ... Is there anything else I can help you with?

                Available tools:
                {tools}

                Tool names: {tool_names}

                Previous conversation:
                {chat_history}

                Human: {input}

                Thought: {agent_scratchpad}
                """)

        agent = create_react_agent(self.llm, tools, prompt)
        self.agent_executor = AgentExecutor(
            agent=agent,
            tools=tools,
            memory=self.memory,
            verbose=True,
            handle_parsing_errors=True,
            max_iterations=8,
            max_execution_time=120
        )

    def add_document(self, file, category: str, description: str, status: str, access: str):
        """Add document to vector store"""
        if self.vectorstore is None or self.embeddings is None:
            st.error("Chatbot components (vector store or embeddings) are not initialized. Please initialize the chatbot first.")
            return False

        tmp_file_path = None # Initialize to None for cleanup in finally block
        try:
            with st.status(f"Processing document: **{file.name}**", expanded=True) as doc_status:
                doc_status.write("Saving uploaded file temporarily...")
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file.name.split('.')[-1]}") as tmp_file:
                    tmp_file.write(file.getvalue())
                    tmp_file_path = tmp_file.name
                doc_status.write("✅ File saved temporarily.")

                file_type = file.name.split('.')[-1].lower()
                doc_status.write(f"Loading document content (type: {file_type})...")
                texts = self.document_processor.load_document(tmp_file_path, file_type)
                if not texts:
                    doc_status.error("Could not extract text from document!")
                    doc_status.update(label=f"Failed to process {file.name}", state="error", expanded=True)
                    return False
                doc_status.write("✅ Document content loaded.")

                metadata = {
                    "filename": file.name,
                    "category": category,
                    "description": description,
                    "status": status,
                    "access": access,
                    "document_id": str(uuid.uuid4())
                }

                doc_status.write("Chunking document into smaller pieces...")
                documents = self.document_processor.chunk_documents(texts, metadata)
                if not documents:
                    doc_status.error("No chunks created from document!")
                    doc_status.update(label=f"Failed to process {file.name}", state="error", expanded=True)
                    return False
                doc_status.write(f"✅ Created {len(documents)} chunks.")

                doc_status.write("Adding chunks to vector database...")
                self.vectorstore.add_documents(documents)
                self.vectorstore.persist()
                doc_status.write("✅ Chunks added to database.")

                doc_status.update(label=f"Successfully added {file.name} ({len(documents)} chunks)", state="complete", expanded=False)
                return True

        except Exception as e:
            # st.error(f"Error adding document: {str(e)}") # st.status takes care of visible error
            doc_status.error(f"An error occurred: {str(e)}")
            doc_status.update(label=f"Failed to add {file.name}", state="error", expanded=True)
            return False
        finally:
            # Ensure temporary file is always cleaned up
            if tmp_file_path and os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)


    def chat(self, message: str) -> str:
        """Process chat message"""
        if not self.agent_executor:
            return "Please initialize the chatbot first!"

        try:
            response = self.agent_executor.invoke({"input": message})
            return response.get("output", "Sorry, I couldn't process your request.")
        except Exception as e:
            return f"Error processing message: {str(e)}"


def main():
    # --- Custom Header Section - NEW ---
    st.markdown('<div class="header-container">', unsafe_allow_html=True)

    st.markdown('<hr style="border: 0.5px solid #eee; margin: 10px 0;">', unsafe_allow_html=True)
    # --- END Custom Header Section ---


    # Initialize session state
    if 'chatbot' not in st.session_state:
        st.session_state.chatbot = RAGChatbot()
    if 'initialized' not in st.session_state:
        st.session_state.initialized = False
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'current_user_type' not in st.session_state:
        st.session_state.current_user_type = "non-support"
    if 'doc_manager' not in st.session_state:
        st.session_state.doc_manager = DocumentManager()
    if 'pending_delete_doc_id' not in st.session_state:
        st.session_state.pending_delete_doc_id = None
    # NEW SESSION STATE VARIABLES FOR MANAGING RENDERING
    if 'processing_message' not in st.session_state:
        st.session_state.processing_message = False # True when bot is thinking
    if 'current_user_input' not in st.session_state: # Renamed from last_user_message for clarity
        st.session_state.current_user_input = None # Stores the current input that triggered processing

    with st.sidebar:
        st.header("Configuration")

        google_api_key = st.text_input("Google API Key", type="password")
        user_type = st.selectbox("User Type", ["support", "non-support"])

        if google_api_key:
            if not st.session_state.initialized:
                if st.button("🚀 Initialize Chatbot", type="primary"):
                    if st.session_state.chatbot.initialize_components(google_api_key):
                        st.session_state.chatbot.setup_agent(user_type)
                        st.session_state.initialized = True
                        st.session_state.current_user_type = user_type
                        st.balloons()
                        st.rerun()
            else:
                if user_type != st.session_state.current_user_type:
                    if st.button("🔄 Update User Type"):
                        if st.session_state.chatbot.llm:
                            st.session_state.chatbot.setup_agent(user_type)
                            st.session_state.current_user_type = user_type
                            st.success(f"Updated to **{user_type}** user!")
                        else:
                            st.error("LLM not initialized. Please re-initialize chatbot.")
                            st.session_state.initialized = False
                            st.rerun()

                if st.button("🔄 Re-initialize Chatbot"):
                    st.session_state.initialized = False
                    st.session_state.chatbot = RAGChatbot()
                    st.session_state.chat_history = []
                    st.session_state.chatbot.memory.clear()
                    st.session_state.doc_manager = DocumentManager()
                    # Reset new session state flags as well
                    st.session_state.processing_message = False
                    st.session_state.current_user_input = None # Reset new flag
                    st.rerun()
        else:
            st.warning("⚠️ Please enter your Google API Key first!")

        if st.session_state.initialized:
            st.success("✅ Chatbot Ready!")
            st.info(f"👤 Current user: **{st.session_state.current_user_type}**")

            if st.session_state.chatbot.vectorstore:
                st.success("🗃️ Vector Database: Connected")
                st.subheader("Database Stats")
                try:
                    count = st.session_state.chatbot.vectorstore._collection.count()
                    st.metric("Total Document Chunks", count)
                    if count > 0:
                        st.success("📚 Documents ready for queries!")
                    else:
                        st.info("📝 Upload documents to get started!")
                except Exception as e:
                    st.warning(f"Stats unavailable: {str(e)}")
            else:
               st.error("🗃️ Vector Database: Not Connected")
               st.error("Please re-initialize the chatbot!")
        else:
            st.info("🤖 Chatbot not initialized")
            st.info("Enter your Google API key and click 'Initialize Chatbot'")

    tab2, tab3 = st.tabs(["📄 Document Upload", "📋 Document Management"])



    with tab2:
        st.header("📄 Document Upload")

        if not st.session_state.initialized:
            st.warning("⚠️ Please initialize the chatbot first!")
            st.info("Steps:\n1. Enter Google API Key in sidebar\n2. Select user type\n3. Click 'Initialize Chatbot'")
        elif st.session_state.chatbot.vectorstore is None:
            st.error("❌ Vector database not available!")
            st.info("Please re-initialize the chatbot using the sidebar.")
        else:
            with st.expander("🔍 Debug Information"):
                st.write("Vectorstore object initialized:", st.session_state.chatbot.vectorstore is not None)
                st.write("Embeddings object initialized:", st.session_state.chatbot.embeddings is not None)
                try:
                    count = st.session_state.chatbot.vectorstore._collection.count()
                    st.write("Current document chunk count:", count)
                except Exception as e:
                    st.write("Error getting count:", str(e))

            with st.form("document_upload"):
                uploaded_file = st.file_uploader(
                    "Choose a document",
                    type=['pdf', 'docx', 'txt'],
                    help="Supported formats: PDF, DOCX, TXT (max 200MB)"
                )

                if uploaded_file:
                    st.info(f"📁 Selected: **{uploaded_file.name}** ({uploaded_file.size / 1024:.2f} KB)")

                col1, col2 = st.columns(2)
                with col1:
                    category = st.selectbox("📂 Category", CATEGORIES, help="Document category")
                    status = st.selectbox("🔄 Status", STATUS_OPTIONS, help="Active documents are searchable")
                with col2:
                    access = st.selectbox("🔐 Access Level", ACCESS_OPTIONS,
                                          help="Internal: Support users only, External: All users")
                    description = st.text_area("📝 Description",
                                               placeholder="Brief description of the document content...",
                                               help="This helps with document discovery")

                submitted = st.form_submit_button("🚀 Upload Document", type="primary")

                if submitted:
                    if uploaded_file is not None:
                        if description.strip():
                            success = st.session_state.chatbot.add_document(
                                uploaded_file, category, description, status, access
                            )
                            if success:
                                st.session_state.doc_manager = DocumentManager()
                                st.rerun()
                        else:
                            st.error("❌ Please provide a description for the document.")
                    else:
                        st.error("❌ Please select a file to upload.")

            st.subheader("📊 Database Overview")
            try:
                collection = st.session_state.chatbot.vectorstore._collection
                count = collection.count()

                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("📚 Total Chunks", count)
                with col2:
                    st.metric("🗃️ Database Type", "ChromaDB")
                with col3:
                    st.metric("👤 Access Context", st.session_state.current_user_type)

                if count > 0:
                    st.success("🎉 Documents are ready for querying!")
                    st.info("💡 Switch to the **Chat** tab to start asking questions about your documents.")
                else:
                    st.info("📁 No documents uploaded yet. Upload your first document above!")

            except Exception as e:
                st.error(f"❌ Error accessing database: {str(e)}")
                st.info("Try re-initializing the chatbot from the sidebar.")
    with tab3:
        st.header("📋 Manage Documents")

        if not st.session_state.initialized:
            st.warning("⚠️ Please initialize the chatbot first to manage documents.")
            st.info("Steps:\n1. Enter Google API Key in sidebar\n2. Select user type\n3. Click 'Initialize Chatbot'")
        elif st.session_state.chatbot.vectorstore is None:
            st.error("❌ Vector database not available! Please re-initialize the chatbot.")
        else:
            all_docs = st.session_state.doc_manager.get_all_documents_metadata()

            if not all_docs:
                st.info("No documents found in the database yet. Upload some in the 'Document Upload' tab.")
            else:
                import pandas as pd
                docs_df = pd.DataFrame(all_docs)

                docs_df['Delete?'] = False

                column_config = {
                    "document_id": st.column_config.Column("Document ID", help="Unique ID of the document",
                                                           disabled=True),
                    "filename": st.column_config.Column("File Name", disabled=True),
                    "description": st.column_config.Column("Description", disabled=True),
                    "total_chunks": st.column_config.NumberColumn("Chunks", disabled=True),
                    "access": st.column_config.SelectboxColumn(
                        "Access Level",
                        help="Change access for this document",
                        options=ACCESS_OPTIONS,
                        required=True
                    ),
                    "status": st.column_config.SelectboxColumn(
                        "Status",
                        help="Change status for this document (Active/Inactive)",
                        options=STATUS_OPTIONS,
                        required=True
                    ),
                    "category": st.column_config.SelectboxColumn(
                        "Category",
                        help="Change category for this document",
                        options=CATEGORIES,
                        required=True
                    ),
                    "Delete?": st.column_config.CheckboxColumn(
                        "Delete?",
                        help="Mark to delete this document",
                        default=False
                    ),
                    "key": st.column_config.Column("Key", width="hidden")
                }

                edited_df = st.data_editor(
                    docs_df,
                    column_config=column_config,
                    hide_index=True,
                    key="doc_data_editor",
                    num_rows="fixed",
                    use_container_width=True
                )

                st.markdown("---")

                if st.button(" Apply Changes ", type="primary"):
                    metadata_changes_applied = False
                    delete_triggered = False

                    for original_row, edited_row in zip(docs_df.to_dict(orient='records'),
                                                        edited_df.to_dict(orient='records')):
                        doc_id_to_process = edited_row['document_id']

                        current_doc_updates = {}
                        if original_row['access'] != edited_row['access']:
                            current_doc_updates["access"] = edited_row['access']
                        if original_row['status'] != edited_row['status']:
                            current_doc_updates["status"] = edited_row['status']
                        if original_row['category'] != edited_row['category']:
                            current_doc_updates["category"] = edited_row['category']

                        if current_doc_updates:
                            with st.spinner(f"Updating fields for {edited_row['filename']}..."):
                                success = st.session_state.doc_manager.update_document_metadata(
                                    doc_id_to_process,
                                    current_doc_updates
                                )
                                if success:
                                    updated_fields_str = ", ".join(
                                        [f"{field}: **{value}**" for field, value in current_doc_updates.items()])
                                    st.success(f"Updated {updated_fields_str} for **{edited_row['filename']}**.")
                                    metadata_changes_applied = True
                                else:
                                    st.error(f"Failed to update fields for **{edited_row['filename']}**.")

                        if edited_row['Delete?'] and not original_row['Delete?']:
                            st.session_state.pending_delete_doc_id = doc_id_to_process
                            delete_triggered = True
                            break

                    if delete_triggered:
                        doc_to_delete_display = next(
                            (d['filename'] for d in st.session_state.doc_manager.get_all_documents_metadata() if
                             d['document_id'] == st.session_state.pending_delete_doc_id), "this document")
                        if st.session_state.doc_manager.delete_document(
                            st.session_state.pending_delete_doc_id):
                            st.success(f"🗑️ Document **{doc_to_delete_display}** deleted successfully.")
                            st.session_state.doc_manager = DocumentManager()
                            st.session_state.pending_delete_doc_id = None
                            st.rerun()
                        else:
                            st.error(f"❌ Failed to delete {doc_to_delete_display}.")
                            st.session_state.pending_delete_doc_id = None
                            st.rerun()


                    elif metadata_changes_applied:
                        st.session_state.doc_manager = DocumentManager()
                        st.rerun()
                    else:
                        st.info("No changes to apply or no documents marked for deletion.")

            st.markdown("---")

            if st.button("🔄 Refresh Document List", key="refresh_doc_list_btn"):
                st.session_state.doc_manager = DocumentManager()
                st.rerun()

if __name__ == "__main__":
    main()